from flask import Flask, request, jsonify
from flask_cors import CORS
from pymongo import MongoClient
from dotenv import load_dotenv
import os
import PyPDF2
from docx import Document
import google.generativeai as genai
from langchain.prompts import ChatPromptTemplate
from langchain.schema import Document as LangchainDocument
from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain.vectorstores import Chroma
from langchain_google_genai import GoogleGenerativeAIEmbeddings

from bson import ObjectId
import io
import re
import logging
from datetime import datetime
import hashlib
import json
from typing import List, Dict, Any, Optional
from werkzeug.utils import secure_filename

# Disable Chroma telemetry
os.environ['ANONYMIZED_TELEMETRY'] = 'False'

app = Flask(__name__)
CORS(app)

# Configure logging
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

# Load environment variables
load_dotenv()
MONGODB_URI = os.getenv("MONGODB_URI", "mongodb://localhost:27017/")
MONGODB_DATABASE = os.getenv("MONGODB_DATABASE", "syllabus_db")
GEMINI_API_KEY = os.getenv("GEMINI_API_KEY")
UPLOAD_FOLDER = os.getenv("UPLOAD_FOLDER", "uploads")
MAX_FILE_SIZE = 16 * 1024 * 1024  # 16MB max file size

# Ensure upload directory exists
os.makedirs(UPLOAD_FOLDER, exist_ok=True)

# MongoDB setup
try:
    client = MongoClient(MONGODB_URI)
    db = client[MONGODB_DATABASE]
    collection = db.syllabus_collections
    # Test connection
    client.admin.command('ping')
    logger.info("Successfully connected to MongoDB")
except Exception as e:
    logger.error(f"Failed to connect to MongoDB: {str(e)}")
    raise

# Initialize AI components
if GEMINI_API_KEY:
    try:        # Configure Gemini
        genai.configure(api_key=GEMINI_API_KEY)
        model = genai.GenerativeModel('gemini-1.5-flash')
        embeddings = GoogleGenerativeAIEmbeddings(model="models/embedding-001", google_api_key=GEMINI_API_KEY)
        
        # Configure text splitting
        text_splitter = RecursiveCharacterTextSplitter(
            chunk_size=1000,
            chunk_overlap=200,
            separators=["\n\n", "\n", ". ", " ", ""]
        )
        
        chat_prompt = """You are an AI assistant specializing in analyzing and answering questions about academic syllabi and course materials. 
        Given the following context from course documents and this question, provide a clear, accurate, and helpful response.
        If the answer cannot be found in the context, say so clearly and suggest what additional information might be helpful.

        Context:
        {context}

        Question:
        {question}

        Answer the question based on the context provided. Be clear and concise, and cite specific sources when relevant.
        """
        
        logger.info("Successfully configured Gemini AI and components")
    except Exception as e:
        logger.error(f"Failed to configure AI components: {str(e)}")
        model = None
        embeddings = None
else:
    logger.warning("GEMINI_API_KEY not found. AI features will be disabled.")
    model = None
    embeddings = None

# Helper function to validate ObjectId
def is_valid_objectid(oid):
    if not oid or oid == 'undefined':
        return False
    return bool(re.match(r'^[0-9a-fA-F]{24}$', oid))

# Helper function to validate file
def allowed_file(filename):
    ALLOWED_EXTENSIONS = {'pdf', 'docx', 'doc', 'txt'}
    return '.' in filename and \
           filename.rsplit('.', 1)[1].lower() in ALLOWED_EXTENSIONS

# Helper function to get file hash
def get_file_hash(file_content):
    return hashlib.md5(file_content).hexdigest()

# Helper function to extract text from PDF
def extract_pdf_text(file):
    try:
        file.seek(0)
        pdf_reader = PyPDF2.PdfReader(file)
        text = ""
        for page_num, page in enumerate(pdf_reader.pages):
            try:
                page_text = page.extract_text()
                if page_text:
                    text += f"\n--- Page {page_num + 1} ---\n{page_text}"
            except Exception as e:
                logger.warning(f"Failed to extract text from page {page_num + 1}: {str(e)}")
                continue
        return text.strip()
    except Exception as e:
        logger.error(f"Failed to extract PDF text: {str(e)}")
        return f"Error extracting PDF: {str(e)}"

# Helper function to extract text from DOCX
def extract_docx_text(file):
    try:
        file.seek(0)
        doc = Document(file)
        text = ""
        for para_num, para in enumerate(doc.paragraphs):
            if para.text.strip():
                text += para.text + "\n"
        
        # Extract text from tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip():
                        text += cell.text + " "
                text += "\n"
        
        return text.strip()
    except Exception as e:
        logger.error(f"Failed to extract DOCX text: {str(e)}")
        return f"Error extracting DOCX: {str(e)}"

# Helper function to extract text from TXT
def extract_txt_text(file):
    try:
        file.seek(0)
        content = file.read()
        if isinstance(content, bytes):
            content = content.decode('utf-8', errors='ignore')
        return content.strip()
    except Exception as e:
        logger.error(f"Failed to extract TXT text: {str(e)}")
        return f"Error extracting TXT: {str(e)}"

# Helper function to extract text based on file type
def extract_file_text(file, filename):
    file_ext = filename.rsplit('.', 1)[1].lower()
    
    if file_ext == 'pdf':
        return extract_pdf_text(file)
    elif file_ext in ['docx', 'doc']:
        return extract_docx_text(file)
    elif file_ext == 'txt':
        return extract_txt_text(file)
    else:
        return "Unsupported file format"

# Error handler
def handle_error(error_msg, status_code=500):
    logger.error(error_msg)
    return jsonify({"error": error_msg, "timestamp": datetime.utcnow().isoformat()}), status_code

# Health check endpoint
@app.route('/api/health', methods=['GET'])
def health_check():
    try:
        # Test MongoDB connection
        client.admin.command('ping')
        mongo_status = "connected"
    except:
        mongo_status = "disconnected"
    
    ai_status = "enabled" if model else "disabled"
    
    return jsonify({
        "status": "healthy",
        "mongodb": mongo_status,
        "ai": ai_status,
        "timestamp": datetime.utcnow().isoformat()
    }), 200

# Get statistics
@app.route('/api/stats', methods=['GET'])
def get_stats():
    try:
        total_collections = collection.count_documents({})
        total_subjects = 0
        total_documents = 0
        
        for coll in collection.find():
            subjects = coll.get('subjects', [])
            total_subjects += len(subjects)
            for subject in subjects:
                total_documents += len(subject.get('documents', []))
        
        return jsonify({
            "total_collections": total_collections,
            "total_subjects": total_subjects,
            "total_documents": total_documents,
            "timestamp": datetime.utcnow().isoformat()
        }), 200
    except Exception as e:
        return handle_error(f"Failed to get statistics: {str(e)}")

# Create a course (collection)
@app.route('/api/collections', methods=['POST'])
def create_collection():
    try:
        data = request.json
        if not data:
            return handle_error("No data provided", 400)
        
        name = data.get('name', '').strip()
        if not name:
            return handle_error("Course name is required", 400)
        
        # Check if collection with same name exists
        existing = collection.find_one({"name": {"$regex": f"^{re.escape(name)}$", "$options": "i"}})
        if existing:
            return handle_error("A course with this name already exists", 409)
        
        course = {
            "name": name,
            "subjects": [],
            "created_at": datetime.utcnow(),
            "updated_at": datetime.utcnow(),
            "description": data.get('description', ''),
            "tags": data.get('tags', [])
        }
        
        result = collection.insert_one(course)
        course['_id'] = str(result.inserted_id)
        
        logger.info(f"Created new course: {name}")
        return jsonify(course), 201
        
    except Exception as e:
        return handle_error(f"Failed to create course: {str(e)}")

# Get all courses (collections) with filtering and sorting
@app.route('/api/collections', methods=['GET'])
def get_collections():
    try:
        # Get query parameters
        search = request.args.get('search', '').strip()
        sort_by = request.args.get('sort_by', 'created_at')
        sort_order = request.args.get('sort_order', 'desc')
        page = int(request.args.get('page', 1))
        limit = int(request.args.get('limit', 50))
        
        # Build query
        query = {}
        if search:
            query["name"] = {"$regex": search, "$options": "i"}
        
        # Build sort
        sort_direction = -1 if sort_order == 'desc' else 1
        sort_query = [(sort_by, sort_direction)]
        
        # Execute query with pagination
        skip = (page - 1) * limit
        courses = list(collection.find(query).sort(sort_query).skip(skip).limit(limit))
        
        # Convert ObjectId to string and add metadata
        for course in courses:
            course['_id'] = str(course['_id'])
            # Add subject and document counts
            course['subject_count'] = len(course.get('subjects', []))
            course['document_count'] = sum(len(s.get('documents', [])) for s in course.get('subjects', []))
        
        # Get total count for pagination
        total_count = collection.count_documents(query)
        
        return jsonify({
            "collections": courses,
            "pagination": {
                "page": page,
                "limit": limit,
                "total": total_count,
                "pages": (total_count + limit - 1) // limit
            }
        }), 200
        
    except Exception as e:
        return handle_error(f"Failed to fetch courses: {str(e)}")

# Get single course by ID
@app.route('/api/collections/<collection_id>', methods=['GET'])
def get_collection(collection_id):
    if not is_valid_objectid(collection_id):
        return handle_error("Invalid course ID", 400)
    
    try:
        course = collection.find_one({"_id": ObjectId(collection_id)})
        if not course:
            return handle_error("Course not found", 404)
        
        course['_id'] = str(course['_id'])
        return jsonify(course), 200
        
    except Exception as e:
        return handle_error(f"Failed to fetch course: {str(e)}")

# Update a course
@app.route('/api/collections/<collection_id>', methods=['PUT'])
def update_collection(collection_id):
    if not is_valid_objectid(collection_id):
        return handle_error("Invalid course ID", 400)
    
    try:
        data = request.json
        if not data:
            return handle_error("No data provided", 400)
        
        update_fields = {}
        if 'name' in data and data['name'].strip():
            update_fields['name'] = data['name'].strip()
        if 'description' in data:
            update_fields['description'] = data['description']
        if 'tags' in data:
            update_fields['tags'] = data['tags']
        
        update_fields['updated_at'] = datetime.utcnow()
        
        result = collection.update_one(
            {"_id": ObjectId(collection_id)},
            {"$set": update_fields}
        )
        
        if result.matched_count == 0:
            return handle_error("Course not found", 404)
        
        return jsonify({"message": "Course updated successfully"}), 200
        
    except Exception as e:
        return handle_error(f"Failed to update course: {str(e)}")

# Delete a course (collection)
@app.route('/api/collections/<collection_id>', methods=['DELETE'])
def delete_collection(collection_id):
    if not is_valid_objectid(collection_id):
        return handle_error("Invalid course ID", 400)
    
    try:
        result = collection.delete_one({"_id": ObjectId(collection_id)})
        if result.deleted_count == 0:
            return handle_error("Course not found", 404)
        
        logger.info(f"Deleted course with ID: {collection_id}")
        return jsonify({"message": "Course deleted successfully"}), 200
        
    except Exception as e:
        return handle_error(f"Failed to delete course: {str(e)}")

# Add a topic (subject) to a course
@app.route('/api/collections/<collection_id>/subjects', methods=['POST'])
def add_subject(collection_id):
    if not is_valid_objectid(collection_id):
        return handle_error("Invalid course ID", 400)
    
    try:
        data = request.json
        if not data:
            return handle_error("No data provided", 400)
        
        name = data.get('name', '').strip()
        if not name:
            return handle_error("Topic name is required", 400)
        
        subject = {
            "name": name,
            "description": data.get('description', ''),
            "documents": [],
            "created_at": datetime.utcnow()
        }
        
        result = collection.update_one(
            {"_id": ObjectId(collection_id)},
            {
                "$push": {"subjects": subject},
                "$set": {"updated_at": datetime.utcnow()}
            }
        )
        
        if result.matched_count == 0:
            return handle_error("Course not found", 404)
        
        logger.info(f"Added subject '{name}' to course {collection_id}")
        return jsonify({"message": "Topic added successfully", "subject": subject}), 201
        
    except Exception as e:
        return handle_error(f"Failed to add topic: {str(e)}")

# Update a subject
@app.route('/api/collections/<collection_id>/subjects/<int:subject_index>', methods=['PUT'])
def update_subject(collection_id, subject_index):
    if not is_valid_objectid(collection_id):
        return handle_error("Invalid course ID", 400)
    
    if subject_index < 0:
        return handle_error("Invalid topic index", 400)
    
    try:
        data = request.json
        if not data:
            return handle_error("No data provided", 400)
        
        update_fields = {}
        if 'name' in data and data['name'].strip():
            update_fields[f'subjects.{subject_index}.name'] = data['name'].strip()
        if 'description' in data:
            update_fields[f'subjects.{subject_index}.description'] = data['description']
        
        update_fields['updated_at'] = datetime.utcnow()
        
        result = collection.update_one(
            {"_id": ObjectId(collection_id), f"subjects.{subject_index}": {"$exists": True}},
            {"$set": update_fields}
        )
        
        if result.matched_count == 0:
            return handle_error("Course or topic not found", 404)
        
        return jsonify({"message": "Topic updated successfully"}), 200
        
    except Exception as e:
        return handle_error(f"Failed to update topic: {str(e)}")

# Delete a topic (subject) from a course
@app.route('/api/collections/<collection_id>/subjects/<int:subject_index>', methods=['DELETE'])
def delete_subject(collection_id, subject_index):
    if not is_valid_objectid(collection_id):
        return handle_error("Invalid course ID", 400)
    
    if subject_index < 0:
        return handle_error("Invalid topic index", 400)
    
    try:
        # First, get the course to check if subject exists
        course = collection.find_one({"_id": ObjectId(collection_id)})
        if not course:
            return handle_error("Course not found", 404)
        
        subjects = course.get('subjects', [])
        if subject_index >= len(subjects):
            return handle_error("Topic not found at the specified index", 404)
        
        # Remove the subject at the specified index
        subjects.pop(subject_index)
        
        result = collection.update_one(
            {"_id": ObjectId(collection_id)},
            {
                "$set": {
                    "subjects": subjects,
                    "updated_at": datetime.utcnow()
                }
            }
        )
        
        if result.matched_count == 0:
            return handle_error("Course not found", 404)
        
        logger.info(f"Deleted subject at index {subject_index} from course {collection_id}")
        return jsonify({"message": "Topic deleted successfully"}), 200
        
    except Exception as e:
        return handle_error(f"Failed to delete topic: {str(e)}")

# Upload syllabus file (document) to a topic
@app.route('/api/collections/<collection_id>/subjects/<int:subject_index>/documents', methods=['POST'])
def upload_document(collection_id, subject_index):
    if not is_valid_objectid(collection_id):
        return handle_error("Invalid course ID", 400)
    
    if subject_index < 0:
        return handle_error("Invalid topic index", 400)
    
    try:
        if 'file' not in request.files:
            return handle_error("No file provided", 400)
        
        file = request.files['file']
        if file.filename == '':
            return handle_error("No file selected", 400)
        
        filename = secure_filename(file.filename)
        if not allowed_file(filename):
            return handle_error("File type not allowed. Only PDF, DOCX, DOC, and TXT files are supported", 400)
        
        # Check file size
        file.seek(0, os.SEEK_END)
        file_size = file.tell()
        if file_size > MAX_FILE_SIZE:
            return handle_error(f"File too large. Maximum size is {MAX_FILE_SIZE // (1024*1024)}MB", 400)
        file.seek(0)
        
        # Extract text from the file
        text_content = extract_file_text(file, filename)
        
        # Get file hash for deduplication
        file.seek(0)
        file_content = file.read()
        file_hash = get_file_hash(file_content)
        
        # Create document object
        document = {
            "filename": filename,
            "content": text_content,
            "file_size": file_size,
            "file_hash": file_hash,
            "uploaded_at": datetime.utcnow(),
            "content_preview": text_content[:500] + "..." if len(text_content) > 500 else text_content
        }
        
        # Check if subject exists and add document
        result = collection.update_one(
            {"_id": ObjectId(collection_id), f"subjects.{subject_index}": {"$exists": True}},
            {
                "$push": {f"subjects.{subject_index}.documents": document},
                "$set": {"updated_at": datetime.utcnow()}
            }
        )
        
        if result.matched_count == 0:
            return handle_error("Course or topic not found", 404)
        
        logger.info(f"Uploaded document '{filename}' to course {collection_id}, subject {subject_index}")
        return jsonify({
            "message": "Document uploaded successfully",
            "filename": filename,
            "file_size": file_size,
            "content_length": len(text_content)
        }), 201
        
    except Exception as e:
        return handle_error(f"Failed to upload document: {str(e)}")

# Delete a syllabus file (document) from a topic
@app.route('/api/collections/<collection_id>/subjects/<int:subject_index>/documents/<int:document_index>', methods=['DELETE'])
def delete_document(collection_id, subject_index, document_index):
    if not is_valid_objectid(collection_id):
        return handle_error("Invalid course ID", 400)
    
    if subject_index < 0 or document_index < 0:
        return handle_error("Invalid topic or document index", 400)
    
    try:
        # First, get the course to check if document exists
        course = collection.find_one({"_id": ObjectId(collection_id)})
        if not course:
            return handle_error("Course not found", 404)
        
        subjects = course.get('subjects', [])
        if subject_index >= len(subjects):
            return handle_error("Topic not found", 404)
        
        documents = subjects[subject_index].get('documents', [])
        if document_index >= len(documents):
            return handle_error("Document not found", 404)
        
        # Remove the document at the specified index
        deleted_doc = documents.pop(document_index)
        subjects[subject_index]['documents'] = documents
        
        result = collection.update_one(
            {"_id": ObjectId(collection_id)},
            {
                "$set": {
                    "subjects": subjects,
                    "updated_at": datetime.utcnow()
                }
            }
        )
        
        if result.matched_count == 0:
            return handle_error("Course not found", 404)
        
        logger.info(f"Deleted document '{deleted_doc.get('filename', 'unknown')}' from course {collection_id}")
        return jsonify({"message": "Document deleted successfully"}), 200
        
    except Exception as e:
        return handle_error(f"Failed to delete document: {str(e)}")

def process_document_for_rag(document_text: str, metadata: Dict[str, Any]) -> List[LangchainDocument]:
    """Process document text into chunks for RAG"""
    chunks = text_splitter.split_text(document_text)
    return [LangchainDocument(page_content=chunk, metadata=metadata) for chunk in chunks]

def get_collection_documents(collection_id: str = None) -> List[LangchainDocument]:
    """Retrieve and process all documents from specified collection(s)"""
    query = {"_id": ObjectId(collection_id)} if collection_id else {}
    
    all_docs = []
    courses = collection.find(query)
    
    for course in courses:
        course_name = course.get('name', 'Unknown Course')
        for subject in course.get('subjects', []):
            subject_name = subject.get('name', 'Unknown Subject')
            for doc in subject.get('documents', []):
                metadata = {
                    'course_name': course_name,
                    'subject_name': subject_name,
                    'filename': doc.get('filename', 'Unknown File'),
                    'source': f"{course_name} > {subject_name} > {doc.get('filename', 'Unknown File')}"
                }
                all_docs.extend(process_document_for_rag(doc.get('content', ''), metadata))
    
    return all_docs



# Enhanced chat endpoint using Gemini directly
@app.route('/api/chat', methods=['POST'])
def chat():
    if not model or not embeddings:
        return handle_error("AI service is not available", 503)
    
    try:
        data = request.json
        if not data:
            return handle_error("No data provided", 400)
        
        query = data.get('query', '').strip()
        collection_id = data.get('collection_id', '').strip()
        
        if not query:
            return handle_error("Query is required", 400)
            
        # Get documents based on collection_id
        docs = get_collection_documents(collection_id if collection_id != 'all' else None)
        
        if not docs:
            return jsonify({
                "answer": "I don't have any documents to work with. Please add some documents first.",
                "sources": []
            }), 200
        
        # Create vector store and get relevant documents
        vectorstore = Chroma.from_documents(docs, embeddings)
        relevant_docs = vectorstore.similarity_search(query, k=5)
        
        # Prepare context from relevant documents
        context = "\n\n".join([
            f"From {doc.metadata.get('source', 'Unknown Source')}:\n{doc.page_content}"
            for doc in relevant_docs
        ])
        
        # Prepare prompt
        prompt = chat_prompt.format(
            context=context,
            question=query
        )
        
        # Generate response using Gemini
        response = model.generate_content(prompt)
        
        # Extract sources
        sources = []
        seen_sources = set()
        for doc in relevant_docs:
            source = doc.metadata.get("source")
            if source and source not in seen_sources:
                seen_sources.add(source)
                sources.append(source)
        
        return jsonify({
            "answer": response.text,
            "sources": sources
        }), 200
        
    except Exception as e:
        logger.error(f"Chat error: {str(e)}")
        return handle_error(f"Failed to process chat request: {str(e)}")


# Search endpoint for cross-course content search
@app.route('/api/search', methods=['POST'])
def search_content():
    try:
        data = request.json
        if not data:
            return handle_error("No data provided", 400)
        
        query = data.get('query', '').strip()
        collection_ids = data.get('collection_ids', [])
        
        if not query:
            return handle_error("Search query is required", 400)
        
        # Build MongoDB query
        search_query = {}
        if collection_ids:
            valid_ids = [ObjectId(cid) for cid in collection_ids if is_valid_objectid(cid)]
            if valid_ids:
                search_query["_id"] = {"$in": valid_ids}
        
        # Search in collection names, subject names, and document content
        courses = list(collection.find(search_query))
        results = []
        
        for course in courses:
            course_matches = []
            
            # Search in course name
            if query.lower() in course.get('name', '').lower():
                course_matches.append({
                    "type": "course_name",
                    "content": course.get('name', ''),
                    "match_type": "title"
                })
            
            # Search in subjects and documents
            for subject_idx, subject in enumerate(course.get('subjects', [])):
                # Search in subject name
                if query.lower() in subject.get('name', '').lower():
                    course_matches.append({
                        "type": "subject_name",
                        "content": subject.get('name', ''),
                        "subject_index": subject_idx,
                        "match_type": "title"
                    })
                
                # Search in document content
                for doc_idx, doc in enumerate(subject.get('documents', [])):
                    doc_content = doc.get('content', '')
                    if query.lower() in doc_content.lower():
                        # Find the context around the match
                        content_lower = doc_content.lower()
                        query_lower = query.lower()
                        match_start = content_lower.find(query_lower)
                        
                        # Extract context (100 chars before and after)
                        context_start = max(0, match_start - 100)
                        context_end = min(len(doc_content), match_start + len(query) + 100)
                        context = doc_content[context_start:context_end]
                        
                        course_matches.append({
                            "type": "document_content",
                            "filename": doc.get('filename', 'Unknown'),
                            "content": context,
                            "subject_index": subject_idx,
                            "document_index": doc_idx,
                            "match_type": "content"
                        })
            
            if course_matches:
                results.append({
                    "course_id": str(course['_id']),
                    "course_name": course.get('name', 'Unknown'),
                    "matches": course_matches[:10]  # Limit matches per course
                })
        
        return jsonify({
            "query": query,
            "results": results,
            "total_courses": len(results),
            "timestamp": datetime.utcnow().isoformat()
        }), 200
        
    except Exception as e:
        return handle_error(f"Failed to search content: {str(e)}")

# Chatbot endpoints for "other" collection
@app.route('/api/chatbot/upload', methods=['POST'])
def upload_chatbot_file():
    """Upload files to the chatbot collection for RAG"""
    try:
        if 'file' not in request.files:
            return handle_error("No file provided", 400)
        
        file = request.files['file']
        if file.filename == '':
            return handle_error("No file selected", 400)
        
        filename = secure_filename(file.filename)
        if not allowed_file(filename):
            return handle_error("File type not allowed. Only PDF, DOCX, DOC, and TXT files are supported", 400)
        
        # Check file size
        file.seek(0, os.SEEK_END)
        file_size = file.tell()
        if file_size > MAX_FILE_SIZE:
            return handle_error(f"File too large. Maximum size is {MAX_FILE_SIZE // (1024*1024)}MB", 400)
        file.seek(0)
        
        # Extract text from the file
        text_content = extract_file_text(file, filename)
        
        # Get file hash for deduplication
        file.seek(0)
        file_content = file.read()
        file_hash = get_file_hash(file_content)
        
        # Check if file already exists
        chatbot_collection = db.chatbot_files
        existing_file = chatbot_collection.find_one({"file_hash": file_hash})
        if existing_file:
            return handle_error("File already exists", 409)
        
        # Create document object
        document = {
            "filename": filename,
            "content": text_content,
            "file_size": file_size,
            "file_hash": file_hash,
            "uploaded_at": datetime.utcnow(),
            "content_preview": text_content[:200] + "..." if len(text_content) > 200 else text_content
        }
        
        result = chatbot_collection.insert_one(document)
        document['_id'] = str(result.inserted_id)
        
        logger.info(f"Uploaded chatbot file: {filename}")
        return jsonify({
            "message": "File uploaded successfully",
            "file": document
        }), 201
        
    except Exception as e:
        return handle_error(f"Failed to upload file: {str(e)}")

@app.route('/api/chatbot/files', methods=['GET'])
def get_chatbot_files():
    """Get all uploaded chatbot files"""
    try:
        chatbot_collection = db.chatbot_files
        files = list(chatbot_collection.find().sort("uploaded_at", -1))
        
        for file in files:
            file['_id'] = str(file['_id'])
            # Remove full content for list view
            file.pop('content', None)
        
        return jsonify({
            "files": files,
            "total": len(files)
        }), 200
        
    except Exception as e:
        return handle_error(f"Failed to get files: {str(e)}")

@app.route('/api/chatbot/files/<file_id>', methods=['DELETE'])
def delete_chatbot_file(file_id):
    """Delete a chatbot file"""
    if not is_valid_objectid(file_id):
        return handle_error("Invalid file ID", 400)
    
    try:
        chatbot_collection = db.chatbot_files
        result = chatbot_collection.delete_one({"_id": ObjectId(file_id)})
        
        if result.deleted_count == 0:
            return handle_error("File not found", 404)
        
        logger.info(f"Deleted chatbot file: {file_id}")
        return jsonify({"message": "File deleted successfully"}), 200
        
    except Exception as e:
        return handle_error(f"Failed to delete file: {str(e)}")

def detect_video_links(text):
    """Detect and extract video links from text"""
    video_patterns = [
        r'https?://(?:www\.)?youtube\.com/watch\?v=[\w-]+',
        r'https?://youtu\.be/[\w-]+',
        r'https?://(?:www\.)?vimeo\.com/\d+',
        r'https?://.*\.(?:mp4|avi|mov|wmv|flv|webm)'
    ]
    
    links = []
    for pattern in video_patterns:
        matches = re.findall(pattern, text, re.IGNORECASE)
        links.extend(matches)
    
    return links

@app.route('/api/chatbot/chat', methods=['POST'])
def chatbot_chat():
    """Enhanced RNS Reply chatbot that uses uploaded files (95%) and general AI knowledge (5%)"""
    if not model:
        return handle_error("AI service is not available", 503)
    
    try:
        data = request.json
        if not data:
            return handle_error("No data provided", 400)
        
        query = data.get('query', '').strip()
        if not query:
            return handle_error("Query is required", 400)
        
        # Get uploaded files for context
        chatbot_collection = db.chatbot_files
        uploaded_files = list(chatbot_collection.find())
        
        context_from_files = ""
        relevant_docs = []
        if uploaded_files and embeddings:
            try:
                # Process uploaded files for RAG
                docs = []
                for file_doc in uploaded_files:
                    metadata = {
                        'filename': file_doc.get('filename', 'Unknown'),
                        'source': f"Uploaded file: {file_doc.get('filename', 'Unknown')}"
                    }
                    docs.extend(process_document_for_rag(file_doc.get('content', ''), metadata))
                
                if docs:
                    # Create vector store and get relevant documents
                    vectorstore = Chroma.from_documents(docs, embeddings)
                    relevant_docs = vectorstore.similarity_search(query, k=5)  # Get more relevant docs
                    
                    context_from_files = "\n\n".join([
                        f"From {doc.metadata.get('source', 'uploaded file')}:\n{doc.page_content}"
                        for doc in relevant_docs
                    ])
            except Exception as e:
                logger.warning(f"Failed to process uploaded files for context: {str(e)}")
          # Create enhanced prompt for RNS Reply with emphasis on uploaded files
        if context_from_files:
            prompt = f"""You are RNS Reply, an advanced AI assistant. Your primary goal is to answer questions using the uploaded documents as your main source of information.

CRITICAL INSTRUCTIONS:
1. **BASE YOUR ANSWER PRIMARILY ON THE UPLOADED FILES (95% of your response)**
2. Use your general knowledge ONLY when the uploaded files don't contain relevant information (maximum 5%)
3. If the uploaded files contain relevant information, prioritize that information heavily
4. When the uploaded files don't have enough information, clearly state this and then provide minimal general knowledge
5. Always cite which uploaded files you're referencing
6. **ALWAYS suggest relevant YouTube videos** that complement your answer, even if files don't mention videos
7. Format YouTube links properly for embedding (use format: https://www.youtube.com/watch?v=VIDEO_ID)

**UPLOADED FILE CONTEXT (THIS IS YOUR PRIMARY SOURCE):**
{context_from_files}

**User Question:** {query}

**Your Response Must:**
- Be based primarily on the uploaded file content above
- Only use general knowledge to fill small gaps if the files don't contain relevant information
- Clearly indicate when you're using information from uploaded files vs general knowledge
- **ALWAYS include 2-3 relevant YouTube video links** that help explain or demonstrate the topic
- Be well-structured and easy to read
- Provide actionable insights

**For YouTube videos:** Even if your uploaded files don't mention specific videos, suggest educational YouTube videos that would help someone learn about this topic. Use real YouTube links in this format: https://www.youtube.com/watch?v=VIDEO_ID

Remember: You are RNS Reply, and your uploaded files are your primary knowledge source, but you should always enhance answers with helpful video resources."""
        else:
            prompt = f"""You are RNS Reply, an advanced AI assistant. Since no files have been uploaded or no relevant information is found in uploaded files, I'll provide a comprehensive response based on general knowledge.

INSTRUCTIONS:
1. Provide detailed, well-researched answers based on general knowledge
2. Be conversational, friendly, and professional
3. **ALWAYS suggest 2-3 relevant YouTube videos** that help explain or demonstrate the topic
4. Format YouTube links properly for embedding (use format: https://www.youtube.com/watch?v=VIDEO_ID)
5. Always aim to be practical and actionable in your advice

**User Question:** {query}

**Note:** No uploaded files contain relevant information for this question.

Provide a comprehensive answer that:
- Is well-researched and accurate based on general knowledge
- **Includes 2-3 relevant YouTube video links** that help explain the topic (format: https://www.youtube.com/watch?v=VIDEO_ID)
- Is well-structured and easy to read
- Offers practical value to the user

**For YouTube videos:** Suggest educational YouTube videos that would help someone learn about this topic. Use real YouTube links in this format: https://www.youtube.com/watch?v=VIDEO_ID

Remember: You are RNS Reply, designed to be your helpful digital assistant with enhanced video recommendations."""
        
        # Generate response using Gemini
        response = model.generate_content(prompt)
        answer = response.text
        
        # Detect video links in the response
        video_links = detect_video_links(answer)
        
        # Prepare sources
        sources = []
        if context_from_files:
            sources = [doc.metadata.get('source') for doc in relevant_docs if doc.metadata.get('source')]
        
        return jsonify({
            "answer": answer,
            "sources": sources,
            "video_links": video_links,
            "has_file_context": bool(context_from_files),
            "chatbot_name": "RNS Reply",
            "timestamp": datetime.utcnow().isoformat()
        }), 200
        
    except Exception as e:
        logger.error(f"RNS Reply chatbot error: {str(e)}")
        return handle_error(f"Failed to process chat request: {str(e)}")

if __name__ == '__main__':
    app.run(debug=True, host='0.0.0.0', port=5000)